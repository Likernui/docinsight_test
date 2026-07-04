"""
Простые тесты базовых правил DocInsight.

Тесты не запускают GUI, GLiNER, PDF и OCR. Они проверяют правила,
которые важны для результата и легко ломаются при доработках.
"""

from src.entity_extractor import EntityExtractor
from src.entity_service import EntityService
from src.preprocessor import TextChunk, TextCleaner, TextPreprocessor
from src.text_extractor import DocxExtractor


def make_extractor() -> EntityExtractor:
    return EntityExtractor(load_model=False)


def test_docx_extractor_preserves_paragraph_table_order(tmp_path):
    print("Проверка: DOCX читается в порядке абзацев и таблиц")
    from docx import Document

    path = tmp_path / "ordered.docx"
    doc = Document()
    doc.add_paragraph("Перед таблицей")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Ячейка 1"
    table.cell(0, 1).text = "Ячейка 2"
    doc.add_paragraph("После таблицы")
    doc.save(path)

    text = DocxExtractor().extract(str(path))

    assert text.splitlines() == [
        "Перед таблицей",
        "Ячейка 1 | Ячейка 2",
        "После таблицы",
    ]


def test_sources_extracted_as_one_block():
    print("Проверка: список литературы извлекается одной сущностью-блоком")
    extractor = make_extractor()
    text = """
Введение

Список использованной литературы
1. Иванов И.И. Основы анализа данных. - Москва, 2020.
2. Петров П.П. Машинное обучение. - Санкт-Петербург, 2021.
"""

    entities = extractor.extract_sources(text)

    assert len(entities) == 1
    assert entities[0]["label"] == "источник"
    assert entities[0]["source_count"] == 2
    assert "Иванов" in entities[0]["text"]
    assert "Петров" in entities[0]["text"]


def test_source_issue_number_is_not_list_item():
    print("Проверка: номер журнала не разбивает источник на два пункта")
    extractor = make_extractor()
    text = """
Список использованной литературы
1. Нажимова Н. А., Нажимов А. В. Автоматизированная система проверки кода // Современные наукоёмкие технологии. 2024. № 6. С. 37–42. URL: https://top-technologies.ru/ru/article/view?id=40061
2. Ковтун М. В. Формирование схемы функциональной структуры [Электронный ресурс]. 2010. URL: https://www.prj-exp.ru/patterns/diagram_functional_structure.php
"""

    entities = extractor.extract_sources(text)

    assert len(entities) == 1
    assert entities[0]["source_count"] == 2
    assert "№ 6. С. 37–42" in entities[0]["text"]
    assert "3. С. 37" not in entities[0]["text"]


def test_sources_with_bibliographic_heading():
    print("Проверка: библиографический список извлекается как источники")
    extractor = make_extractor()
    text = """
Заключение

Библиографический список
1. Иванов И.И. Основы программирования: учебное пособие. - Москва, 2020.
2. Петров П.П. Анализ данных // Научный журнал. - 2021. - N 4.
"""

    entities = extractor.extract_sources(text)

    assert len(entities) == 1
    assert entities[0]["source_count"] == 2
    assert "Основы программирования" in entities[0]["text"]


def test_unnumbered_sources_split_by_paragraphs():
    print("Проверка: ненумерованные источники делятся по абзацам")
    extractor = make_extractor()
    text = """
СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ
OPC Foundation. OPC Unified Architecture. Part 3: Address Space Model. Release 1.05.04 — 2024. — 112 с. Описание: стандарт OPC UA.
ГОСТ 19.201-78. ЕСПД. Техническое задание. Требования к содержанию и оформлению // Разработка технической документации : сайт – URL: https://example.com/gost.pdf
Документация Avalonia UI // Avalonia UI : сайт – URL: https://docs.avaloniaui.net/docs/welcome
ПЛАН-ГРАФИК РАБОТЫ НАД ПРОЕКТОМ
1. Анализ требований С 02.03.2025 по 15.03.2025 Полностью выполнено
"""

    entities = extractor.extract_sources(text)

    assert len(entities) == 1
    assert entities[0]["source_count"] == 3
    assert "OPC Foundation" in entities[0]["text"]
    assert "Техническое задание. Требования" in entities[0]["text"]
    assert "ПЛАН-ГРАФИК" not in entities[0]["text"]


def test_schedule_is_not_source():
    print("Проверка: план-график не считается списком источников")
    extractor = make_extractor()
    text = """
План-график выполнения проекта
1. Постановка задачи.
2. Разработка интерфейса.
3. Подготовка отчета.
"""

    entities = extractor.extract_sources(text, bottom_section=True)

    assert entities == []


def test_schedule_with_dates_is_not_source():
    print("Проверка: план-график с датами не считается списком источников")
    extractor = make_extractor()
    text = """
1. Определение архитектуры проекта С 02.03.2025 по 15.03.2025 Растопчин П. Полностью выполнено
2. Разработка главной страницы и навигации С 16.03.2025 по 23.03.2025 Ефимов А.Д. Полностью выполнено
3. Разработка модуля управления сессией OPC UA С 16.03.2025 по 26.03.2025 Растопчин П. Полностью выполнено
"""

    entities = extractor.extract_sources(text, bottom_section=True)

    assert entities == []


def test_inline_source_heading_stops_before_schedule():
    print("Проверка: источники после чанкинга не захватывают план-график")
    extractor = make_extractor()
    text = (
        "Заключение. Список источников "
        "1. Moodle API Documentation [Электронный ресурс]. - URL: https://docs.moodle.org/ "
        "(дата обращения: 10.06.2025). "
        "2. PostgreSQL Documentation [Электронный ресурс]. - URL: https://www.postgresql.org/docs/ "
        "(дата обращения: 10.06.2025). "
        "ПЛАН-ГРАФИК РАБОТЫ НАД ПРОЕКТОМ "
        "1. Анализ требований С 02.03.2025 по 15.03.2025 Полностью выполнено"
    )

    entities = extractor.extract_sources(text)

    assert len(entities) == 1
    assert entities[0]["source_count"] == 2
    assert "Moodle API Documentation" in entities[0]["text"]
    assert "PostgreSQL Documentation" in entities[0]["text"]
    assert "ПЛАН-ГРАФИК" not in entities[0]["text"]


def test_source_section_ignores_cover_and_tables():
    print("Проверка: источники не начинаются с обложки и таблиц")
    extractor = make_extractor()
    text = """
Система собирает данные из открытых источников.
Тюмень, 2025

Таблица 2 Атрибуты сущности Server
Атрибут | Тип данных | Описание атрибута
Id | Целое число | Уникальный идентификатор

Список источников
1. OPC Foundation Documentation [Электронный ресурс]. - URL: https://opcfoundation.org/
2. Минин П. Е. Анализ существующих автоматизированных систем управления технологическим процессом // Спецтехника и связь. 2014. №1. URL: https://cyberleninka.ru/article/n/test
"""

    entities = extractor.extract_sources(text, bottom_section=True)

    assert len(entities) == 1
    assert entities[0]["source_count"] == 2
    assert "Тюмень, 2025" not in entities[0]["text"]
    assert "Таблица 2" not in entities[0]["text"]
    assert "OPC Foundation Documentation" in entities[0]["text"]


def test_sources_stop_before_technical_assignment():
    print("Проверка: список литературы не захватывает техническое задание")
    extractor = make_extractor()
    text = """
Список использованной литературы
1. Пухов С. И. Создание системы сбора и обработки открытых данных // Теплотехника и информатика. - Екатеринбург, 2018.
2. Васина Е. Н. Проблема структуризации современных информационных ресурсов // Вестник РЭА. 2014. №4.
ТЕХНИЧЕСКОЕ ЗАДАНИЕ
1.1. Веб-сервис предназначен для автоматизации сбора данных.
"""

    entities = extractor.extract_sources(text)

    assert len(entities) == 1
    assert entities[0]["source_count"] == 2
    assert "ТЕХНИЧЕСКОЕ ЗАДАНИЕ" not in entities[0]["text"]
    assert "Веб-сервис предназначен" not in entities[0]["text"]


def test_sources_are_attached_once_after_chunking():
    print("Проверка: источники прикрепляются один раз после разбиения на чанки")
    text = """
Заключение.

Список источников
1. Moodle API Documentation [Электронный ресурс]. - URL: https://docs.moodle.org/ (дата обращения: 10.06.2025).
2. PostgreSQL Documentation [Электронный ресурс]. - URL: https://www.postgresql.org/docs/ (дата обращения: 10.06.2025).

ПЛАН-ГРАФИК РАБОТЫ НАД ПРОЕКТОМ
1. Анализ требований С 02.03.2025 по 15.03.2025 Полностью выполнено
"""
    chunks = TextPreprocessor(chunk_size=120, overlap_sentences=0).process(
        text,
        "/tmp/report.docx",
    )

    service = EntityService()
    enriched = service.enrich_chunks_regex_only(
        {"/tmp/report.docx": chunks},
        source_texts={"/tmp/report.docx": text},
    )
    rows = service.flatten_entities(enriched)
    sources = [row for row in rows if row["type"] == "источник"]

    assert len(sources) == 1
    assert sources[0]["metadata"]["source_count"] == 2
    assert "ПЛАН-ГРАФИК" not in sources[0]["metadata"]["text"]


def test_old_chunk_source_is_replaced_by_document_source():
    print("Проверка: старый source из чанка заменяется источником из полного документа")
    document_path = "/tmp/report.pdf"
    document_text = """
Список использованной литературы
1. Перязева Ю. В. Возможности автоматической проверки заданий в LMS Moodle // Современные информационные технологии и ИТ-образование. 2019. Т. 15, № 4. С. 876–885.
2. Нажимова Н. А., Нажимов А. В. Автоматизированная система проверки кода // Современные наукоёмкие технологии. 2024. № 6. С. 37–42. URL: https://top-technologies.ru/ru/article/view?id=40061
"""
    chunks = [
        TextChunk(
            text="Старый фрагмент с плохой склейкой",
            file_path=document_path,
            chunk_index=0,
            metadata={
                "entities": [
                    {
                        "label": "источник",
                        "text": (
                            "Список использованной литературы\n"
                            "1. Перязева Ю. В. Возможности автоматической проверки заданий. "
                            "По результатам проверки будет формироваться отчёт"
                        ),
                        "title": "Список использованной литературы",
                        "score": 1.0,
                        "source": "regex",
                        "source_count": 1,
                    }
                ]
            },
        ),
        TextChunk(
            text=document_text,
            file_path=document_path,
            chunk_index=1,
            metadata={},
        ),
    ]

    service = EntityService()
    enriched = service.enrich_chunks_regex_only(
        {document_path: chunks},
        source_texts={document_path: document_text},
    )
    rows = service.flatten_entities(enriched)
    sources = [row for row in rows if row["type"] == "источник"]

    assert len(sources) == 1
    assert sources[0]["metadata"]["source_count"] == 2
    assert "По результатам проверки" not in sources[0]["metadata"]["text"]
    assert "№ 6. С. 37–42" in sources[0]["metadata"]["text"]


def test_code_block_extracted_as_one_entity():
    print("Проверка: функция извлекается одним кодовым блоком")
    extractor = make_extractor()
    text = """
Ниже приведен фрагмент программы:

def assess_answer(answer):
    score = len(answer)
    return score
"""

    entities = extractor.extract_regex_entities(text, include_sources=False)
    code_entities = [
        entity
        for entity in entities
        if entity["label"] == "фрагмент программного кода"
    ]

    assert len(code_entities) == 1
    assert code_entities[0]["title"] == "функция assess_answer"
    assert "return score" in code_entities[0]["code"]


def test_single_return_is_ignored():
    print("Проверка: одиночный return не считается кодовым фрагментом")
    extractor = make_extractor()
    text = "Алгоритм завершает работу и возвращает результат.\nreturn score"

    entities = extractor.extract_regex_entities(text, include_sources=False)
    code_entities = [
        entity
        for entity in entities
        if entity["label"] == "фрагмент программного кода"
    ]

    assert code_entities == []


def test_gliner_rejects_label_like_person_name():
    print("Проверка: GLiNER-ошибка 'ФИО студентов' не считается человеком")
    extractor = make_extractor()

    entities = extractor._normalize_entities(
        "В таблице указаны ФИО студентов.",
        [
            {
                "text": "ФИО студентов",
                "label": "ФИО человека",
                "score": 0.95,
                "start": 17,
                "end": 30,
            }
        ],
    )

    assert entities == []


def test_students_extracted_from_labeled_block():
    print("Проверка: студенты извлекаются из блока с маркером")
    extractor = make_extractor()
    text = """
Выполнили работу студенты
Кепещук Марьяна Денисовна
Жакаев Даулет Канатович
"""

    entities = extractor.extract_regex_entities(text, include_sources=False)
    students = [
        entity["text"]
        for entity in entities
        if entity["label"] == "ФИО студента"
    ]

    assert "Кепещук Марьяна Денисовна" in students
    assert "Жакаев Даулет Канатович" in students


def test_title_page_students_and_supervisor_are_not_mixed():
    print("Проверка: студенты и руководитель из титульника не смешиваются")
    extractor = make_extractor()
    text = """
Выполнили работу
студенты 3 курса
очной формы обучения
Ершов М.М. Жилин С.А. Руководитель проекта
Перевалова М.Н.
"""

    entities = extractor.extract_regex_entities(text, include_sources=False)
    students = {
        entity["text"]
        for entity in entities
        if entity["label"] == "ФИО студента"
    }
    supervisors = {
        entity["text"]
        for entity in entities
        if entity["label"] == "ФИО куратора"
    }

    assert "Ершов М.М." in students
    assert "Жилин С.А." in students
    assert "Перевалова М.Н." in supervisors
    assert "Ершов М.М." not in supervisors
    assert "Жилин С.А." not in supervisors


def test_source_author_is_not_person_entity_without_role_marker():
    print("Проверка: автор источника без роли не считается студентом или руководителем")
    extractor = make_extractor()
    text = """
Список использованной литературы
1. Иванов Иван Иванович. Основы анализа данных. - Москва, 2020.
2. Петров Петр Петрович. Анализ систем // Научный журнал. 2021. URL: https://example.com
Описание: Автор рассматривает современные подходы.
"""

    entities = extractor.extract_regex_entities(text, include_sources=False)
    person_entities = [
        entity
        for entity in entities
        if entity["label"] in {"ФИО студента", "ФИО куратора", "ФИО"}
    ]

    assert person_entities == []


def test_supervisor_extracted_from_labeled_block():
    print("Проверка: руководитель извлекается из подписи")
    extractor = make_extractor()
    text = """
Выполнил студент Иванов И.И.

Руководитель проекта
Мельникова Антонина Владимировна
"""

    entities = extractor.extract_regex_entities(text, include_sources=False)
    supervisors = [
        entity
        for entity in entities
        if entity["label"] == "ФИО куратора"
    ]

    assert len(supervisors) == 1
    assert supervisors[0]["text"] == "Мельникова Антонина Владимировна"


def test_title_persons_extracted_across_first_chunks():
    print("Проверка: ФИО на титульнике извлекаются через границу чанков")
    extractor = make_extractor()
    chunks = [
        TextChunk(
            text="Итоговый отчет\nВыполнили работу студенты",
            file_path="report.pdf",
            chunk_index=0,
        ),
        TextChunk(
            text="Дронов Илья Сергеевич\nЗимин Андрей Валерьевич______ Лишний Куратор\nРуководитель",
            file_path="report.pdf",
            chunk_index=1,
        ),
        TextChunk(
            text="Воробьева\nМарина\nСергеевна\nТюмень, 2025",
            file_path="report.pdf",
            chunk_index=2,
        ),
    ]

    extractor.attach_title_person_entities(chunks)
    entities = chunks[0].metadata["entities"]
    students = {
        entity["text"]
        for entity in entities
        if entity["label"] == "ФИО студента"
    }
    supervisors = {
        entity["text"]
        for entity in entities
        if entity["label"] == "ФИО куратора"
    }

    assert students == {"Дронов Илья Сергеевич", "Зимин Андрей Валерьевич"}
    assert supervisors == {"Воробьева Марина Сергеевна"}


def test_title_table_merged_names_are_split_by_role():
    print("Проверка: склеенная таблица титульника делит студентов и руководителя")
    extractor = make_extractor()
    text = """
Выполнили работу
(групповой проект) студенты 3 курса очной формы обучения Руководитель | Ефимов Артем Денисович Растопчин Павел
Аврискин Михаил Владимирович
Тюмень, 2025
"""

    entities = extractor.extract_title_person_fields(text)
    students = {
        entity["text"]
        for entity in entities
        if entity["label"] == "ФИО студента"
    }
    supervisors = {
        entity["text"]
        for entity in entities
        if entity["label"] == "ФИО куратора"
    }

    assert students == {"Ефимов Артем Денисович", "Растопчин Павел"}
    assert supervisors == {"Аврискин Михаил Владимирович"}


def test_quality_audio_is_not_person_name():
    print("Проверка: названия показателей не считаются ФИО")
    extractor = make_extractor()
    text = """
Было опрошено 5 студентов и 1 преподаватель.
№ | Качество Аудио (1-10) | Качество Видео (1-10)
1 | 7 | 10
"""

    entities = extractor.extract_regex_entities(text, include_sources=False)
    person_entities = [
        entity
        for entity in entities
        if entity["label"] in {"ФИО студента", "ФИО куратора"}
    ]

    assert person_entities == []


def test_person_sources_are_combined():
    print("Проверка: источники ФИО объединяются при совпадении regex и Natasha")
    extractor = make_extractor()

    entities = extractor.deduplicate_entities([
        {
            "text": "Иванов Иван Иванович",
            "label": "ФИО студента",
            "score": 1.0,
            "source": "regex",
        },
        {
            "text": "Иванов Иван Иванович",
            "label": "ФИО студента",
            "score": 0.95,
            "source": "natasha",
        },
    ])

    assert len(entities) == 1
    assert entities[0]["source"] == "regex+natasha"


def test_word_indented_code_block_is_preserved():
    print("Проверка: код с Word-отступами извлекается целым блоком")
    extractor = make_extractor()
    text = (
        "def create(self, request):\n"
        "\u00a0\u00a0\u00a0\u00a0items = CartItem.objects.all()\n"
        "\u00a0\u00a0\u00a0\u00a0for cart_item in items:\n"
        "\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0# Проверяем доступность\n"
        "\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0if cart_item.itemQuantity > 0:\n"
        "\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0return Response({\n"
        "\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\"status\": \"success\",\n"
        "\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0\u00a0})\n"
    )
    cleaned_text = TextCleaner.clean(text)
    cleaned = extractor.extract_regex_entities(cleaned_text, include_sources=False)
    code_entities = [
        entity
        for entity in cleaned
        if entity["label"] == "фрагмент программного кода"
    ]

    assert len(code_entities) == 1
    assert "for cart_item in items" in code_entities[0]["code"]
    assert "# Проверяем доступность" in code_entities[0]["code"]
    assert '"status": "success"' in code_entities[0]["code"]


def test_flatten_entities_for_ui():
    print("Проверка: сущности преобразуются в строки для пользовательского UI")
    chunk = TextChunk(
        text="def assess_answer(answer):\n    return len(answer)",
        file_path="/tmp/report.docx",
        chunk_index=0,
        metadata={
            "entities": [
                {
                    "label": "фрагмент программного кода",
                    "text": "функция assess_answer",
                    "title": "функция assess_answer",
                    "code": "def assess_answer(answer):\n    return len(answer)",
                    "score": 1.0,
                    "source": "regex",
                    "chunk_indexes": [0],
                }
            ]
        },
    )

    rows = EntityService().flatten_entities({chunk.file_path: [chunk]})

    assert len(rows) == 1
    assert rows[0]["type"] == "фрагмент программного кода"
    assert rows[0]["value"] == "функция assess_answer"
    assert rows[0]["document"] == "report.docx"
    assert rows[0]["chunk_indexes"] == [0]
    assert rows[0]["confidence"] == 1.0
    assert rows[0]["source"] == "regex"
    assert "code" in rows[0]["metadata"]
