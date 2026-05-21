"""
Генерация простого DOCX-отчета по найденным сущностям.
"""

from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


ENTITY_SECTION_TITLES = {
    "ФИО студента": "Студенты и участники",
    "ФИО куратора": "Руководители",
    "ФИО": "ФИО",
    "ответственный": "Ответственные",
    "тема проекта": "Темы проектов",
    "тема работы": "Темы работ",
    "название проекта": "Названия проектов",
    "дисциплина": "Дисциплины",
    "технология": "Технологии",
    "источник": "Источники",
    "фрагмент программного кода": "Фрагменты программного кода",
}
ENTITY_SECTION_ORDER = {
    entity_type: index
    for index, entity_type in enumerate(ENTITY_SECTION_TITLES)
}


class DocxReportGenerator:
    def generate(
        self,
        rows: list[dict[str, Any]],
        output_path: str | Path,
        title: str = "Отчет DocInsight",
        scope: str = "Все документы",
    ) -> None:
        document = Document()
        document.add_heading(title, level=0)
        document.add_paragraph(f"Область отчета: {scope}")
        document.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        document.add_paragraph(f"Найдено строк: {len(rows)}")

        grouped = self._group_rows(rows)
        if not grouped:
            document.add_paragraph("Сущности не найдены.")
            document.save(output_path)
            return

        for entity_type in sorted(grouped.keys(), key=self._section_sort_key):
            document.add_heading(ENTITY_SECTION_TITLES.get(entity_type, entity_type), level=1)

            for row in grouped[entity_type]:
                self._add_row(document, row)

        document.save(output_path)

    def _group_rows(self, rows: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
        grouped = defaultdict(list)
        for row in rows:
            grouped[str(row.get("type") or "неизвестно")].append(row)
        return dict(grouped)

    def _section_sort_key(self, entity_type: str) -> tuple[int, str]:
        return ENTITY_SECTION_ORDER.get(entity_type, len(ENTITY_SECTION_ORDER)), entity_type

    def _add_row(self, document: Document, row: dict[str, Any]) -> None:
        value = str(row.get("value") or "").strip()
        metadata = row.get("metadata") or {}

        if row.get("type") == "фрагмент программного кода":
            document.add_paragraph(value, style="List Bullet")
            code = str(metadata.get("code") or metadata.get("text") or "").strip()
            if code:
                document.add_paragraph(code)
            return

        if row.get("type") == "источник":
            document.add_paragraph(value, style="List Bullet")
            text = str(metadata.get("text") or "").strip()
            if text and text != value:
                document.add_paragraph(text)
            return

        details = self._details_text(row)
        paragraph_text = value if not details else f"{value} ({details})"
        document.add_paragraph(paragraph_text, style="List Bullet")

    def _details_text(self, row: dict[str, Any]) -> str:
        parts = []

        document_name = row.get("document")
        if document_name:
            parts.append(str(document_name))

        confidence = row.get("confidence")
        if confidence is not None:
            parts.append(f"уверенность: {float(confidence):.2f}")

        source = row.get("source")
        if source:
            parts.append(f"источник: {source}")

        return ", ".join(parts)
